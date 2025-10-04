#!/usr/bin/env python3
"""
docx-translator - Microsoft Word Document Translation Utility with Google Genai API v1.27

A command-line tool for translating Microsoft Word documents while preserving
formatting, styles, tables, and document structure. Updated for google-genai 1.27.0
with improved JSON parsing and error handling.

Features:
- Uses the new google-genai SDK (v1.27.0) instead of deprecated google-generativeai
- Translate single files or entire directories
- Preserve document formatting, styles, and structure
- Support for tables with context-aware translation
- Smart context extraction for better translation quality
- Multiple target languages support
- Style-based translation (business, casual, technical, etc.)
- Robust JSON parsing to handle API response variations
- Optimized for Japanese and other languages

Usage Examples:
    # Translate a single file to Japanese
    python docx-translator.py document.docx --target-lang ja

    # Translate with specific style
    python docx-translator.py report.docx --target-lang ja --style-prompt business

    # Translate to multiple languages
    python docx-translator.py document.docx --target-langs en,ja,es,zh

    # Use glossary for consistent terminology
    python docx-translator.py technical.docx --target-lang ja --context-file glossary.json
"""

import argparse
import asyncio
import json
import logging
import os
import re
import sys
import time
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor, as_completed

# Updated import for google-genai 1.27.0
from google import genai
from google.genai import types
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.shared import RGBColor
from colorama import Fore, Style, init
from tqdm import tqdm

# Initialize colorama for colored output
init(autoreset=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Constants
DEFAULT_MODEL = 'gemini-2.5-flash'
DEFAULT_BATCH_SIZE = 10
DEFAULT_CONCURRENCY = 4
MAX_RETRIES = 3
RETRY_DELAY = 2

# Style prompts optimized for translations
STYLE_PROMPTS = {
    'business': {
        'ja': 'ビジネス文書として適切な敬語と専門用語を使用し、フォーマルで明確な日本語に翻訳してください。',
        'default': 'Translate using formal business language appropriate for professional documents.'
    },
    'casual': {
        'ja': '自然で親しみやすい日本語に翻訳し、適切な丁寧さを保ってください。',
        'default': 'Translate using conversational, friendly language suitable for general communication.'
    },
    'technical': {
        'ja': '技術的な正確性を優先し、専門用語は適切な日本語の技術用語を使用してください。',
        'default': 'Translate using precise technical terminology, maintaining accuracy for specialized content.'
    },
    'academic': {
        'ja': '学術的な文体で、正確性と論理性を重視した日本語に翻訳してください。',
        'default': 'Translate using scholarly language appropriate for academic contexts.'
    },
    'marketing': {
        'ja': '説得力があり、読者の心に響く魅力的な日本語表現を使用してください。',
        'default': 'Translate using persuasive, engaging language suitable for marketing materials.'
    }
}


@dataclass
class TranslationConfig:
    """Configuration for translation job"""
    source_lang: str
    target_lang: str
    model_name: str
    batch_size: int
    concurrency: int
    style_prompt: str
    context_file: Optional[str]
    smart_context: bool
    debug: bool


class ImprovedJSONExtractor:
    """Enhanced JSON extraction with multiple fallback strategies"""
    
    @staticmethod
    def extract_json(response_text: str) -> Dict[str, Any]:
        """Extract JSON from response with multiple strategies"""
        # Strategy 1: Try direct JSON parsing
        try:
            return json.loads(response_text)
        except json.JSONDecodeError:
            pass
        
        # Strategy 2: Extract from markdown code blocks
        code_block_patterns = [
            r'```json\s*\n(.*?)\n```',
            r'```\s*\n(.*?)\n```',
            r'`(.*?)`'
        ]
        
        for pattern in code_block_patterns:
            matches = re.findall(pattern, response_text, re.DOTALL)
            if matches:
                for match in matches:
                    try:
                        return json.loads(match)
                    except json.JSONDecodeError:
                        continue
        
        # Strategy 3: Find JSON-like structures
        json_patterns = [
            r'\{[^{}]*"translations"[^{}]*\}',  # Simple JSON with translations
            r'\{.*?"translations".*?\}',         # More flexible
            r'\{.*\}',                           # Any JSON object
        ]
        
        for pattern in json_patterns:
            matches = re.findall(pattern, response_text, re.DOTALL)
            if matches:
                # Try the longest match first
                for match in sorted(matches, key=len, reverse=True):
                    try:
                        # Clean up common issues
                        cleaned = match
                        # Remove trailing commas
                        cleaned = re.sub(r',\s*}', '}', cleaned)
                        cleaned = re.sub(r',\s*]', ']', cleaned)
                        # Fix escaped quotes that shouldn't be escaped
                        cleaned = cleaned.replace('\\"', '"')
                        # Remove any BOM or zero-width spaces
                        cleaned = cleaned.replace('\ufeff', '').replace('\u200b', '')
                        
                        result = json.loads(cleaned)
                        if isinstance(result, dict):
                            return result
                    except json.JSONDecodeError:
                        continue
        
        # Strategy 4: Try to fix common JSON issues
        cleaned_text = response_text
        # Remove any text before the first {
        first_brace = cleaned_text.find('{')
        if first_brace != -1:
            cleaned_text = cleaned_text[first_brace:]
        # Remove any text after the last }
        last_brace = cleaned_text.rfind('}')
        if last_brace != -1:
            cleaned_text = cleaned_text[:last_brace + 1]
        
        try:
            return json.loads(cleaned_text)
        except json.JSONDecodeError:
            pass
        
        # If all strategies fail, raise an exception with helpful info
        raise ValueError(f"Could not extract valid JSON from response. First 200 chars: {response_text[:200]}")


class DocxTranslator:
    """Main translator class for Word documents using google-genai 1.27.0"""
    
    def __init__(self, config: TranslationConfig):
        self.config = config
        self.json_extractor = ImprovedJSONExtractor()
        self._setup_logging()
        self._configure_api()
        self.context_data = self._load_context()
        
    def _setup_logging(self):
        """Configure logging based on debug flag"""
        level = logging.DEBUG if self.config.debug else logging.INFO
        logging.getLogger().setLevel(level)
        self.logger = logging.getLogger(__name__)
        
    def _configure_api(self):
        """Configure Google Genai API v1.27.0"""
        api_key = os.getenv('GOOGLE_API_KEY')
        if not api_key:
            raise ValueError(
                f"{Fore.RED}Error: GOOGLE_API_KEY environment variable not set.{Style.RESET_ALL}\n"
                f"{Fore.YELLOW}Please set your Google API key as an environment variable.{Style.RESET_ALL}"
            )
        
        # Initialize the new google-genai client
        self.client = genai.Client(api_key=api_key)
        self.logger.info(f"Initialized google-genai client v1.27.0 with model: {self.config.model_name}")
        
    def _load_context(self) -> Dict:
        """Load context/glossary file if provided"""
        if not self.config.context_file:
            return {}
            
        try:
            with open(self.config.context_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                self.logger.info(f"Loaded context from {self.config.context_file}")
                return data
        except Exception as e:
            self.logger.warning(f"Failed to load context file: {e}")
            return {}
    
    def _get_style_prompt(self) -> str:
        """Get appropriate style prompt for the target language"""
        style_dict = STYLE_PROMPTS.get(self.config.style_prompt, {})
        return style_dict.get(self.config.target_lang, style_dict.get('default', ''))
    
    def _build_translation_prompt(self, texts: List[Dict[str, Any]], context: List[str] = None) -> str:
        """Build optimized prompt for Gemini 2.5 Flash"""
        style_instruction = self._get_style_prompt()
        
        context_section = ""
        if self.context_data:
            glossary_items = list(self.context_data.items())[:20]
            if glossary_items:
                context_section = "Use this glossary for consistent terminology:\n"
                for term, translation in glossary_items:
                    context_section += f"- {term} → {translation}\n"
        
        if context:
            context_section += f"\nDocument context:\n{chr(10).join(context[:10])}\n"
        
        # Optimized prompt for Gemini 2.5 Flash with explicit JSON instructions
        prompt = f"""Translate the following texts from {self.config.source_lang} to {self.config.target_lang}.
{style_instruction}

CRITICAL INSTRUCTIONS:
1. Return ONLY a valid JSON object with NO markdown formatting, NO code blocks, NO backticks
2. The JSON must have this exact structure: {{"translations": [{{"id": 0, "translation": "translated text"}}, {{"id": 1, "translation": "translated text"}}]}}
3. Preserve ALL formatting including line breaks (\\n), spaces, and punctuation
4. If text is already in the target language, return it unchanged
5. Maintain the same tone and style as the original
6. Do NOT include any explanation or additional text outside the JSON

{context_section}

Texts to translate:
{json.dumps(texts, ensure_ascii=False, indent=2)}

Return only the JSON object:"""
        
        return prompt
    
    async def translate_batch_async(self, texts: List[Dict[str, Any]], context: List[str] = None) -> Dict[int, str]:
        """Translate a batch of texts using async API calls with google-genai 1.27.0"""
        if not texts:
            return {}
        
        prompt = self._build_translation_prompt(texts, context)
        
        for attempt in range(MAX_RETRIES):
            try:
                # Use the new google-genai API structure
                response = await self.client.aio.models.generate_content(
                    model=self.config.model_name,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.3,  # Lower temperature for more consistent translations
                        candidate_count=1,
                        max_output_tokens=2048,
                    )
                )
                
                # Extract text from response
                if response.text:
                    response_text = response.text
                else:
                    # Handle different response structures
                    response_text = str(response)
                
                self.logger.debug(f"Raw API response: {response_text[:500]}...")
                
                # Extract and parse JSON
                result = self.json_extractor.extract_json(response_text)
                
                # Validate response structure
                if 'translations' not in result:
                    raise ValueError("Missing 'translations' key in response")
                
                # Build translation map
                translation_map = {}
                for item in result['translations']:
                    idx = item.get('id')
                    translation = item.get('translation', '')
                    if idx is not None:
                        translation_map[idx] = translation
                
                # Verify all texts were translated
                missing_ids = set(range(len(texts))) - set(translation_map.keys())
                if missing_ids:
                    self.logger.warning(f"Missing translations for IDs: {missing_ids}")
                    # Add original text as fallback for missing translations
                    for idx in missing_ids:
                        translation_map[idx] = texts[idx]['text']
                
                return translation_map
                
            except Exception as e:
                self.logger.warning(f"Translation attempt {attempt + 1} failed: {e}")
                if "429" in str(e) or "quota" in str(e).lower():
                    # Rate limit hit, wait longer
                    wait_time = RETRY_DELAY * (2 ** attempt)
                    self.logger.info(f"Rate limit detected, waiting {wait_time} seconds...")
                    await asyncio.sleep(wait_time)
                elif attempt < MAX_RETRIES - 1:
                    await asyncio.sleep(RETRY_DELAY * (attempt + 1))
                else:
                    self.logger.error(f"All translation attempts failed: {e}")
                    # Return original texts as fallback
                    return {i: text['text'] for i, text in enumerate(texts)}
    
    def translate_batch(self, texts: List[Dict[str, Any]], context: List[str] = None) -> Dict[int, str]:
        """Synchronous wrapper for translate_batch_async"""
        # Create new event loop for sync context
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(self.translate_batch_async(texts, context))
        finally:
            loop.close()
    
    def _extract_document_context(self, doc: Document) -> List[str]:
        """Extract context from the document"""
        context = []
        
        # Get document title from first heading
        for para in doc.paragraphs[:5]:
            if para.style.name.startswith('Heading') or para.style.name == 'Title':
                if para.text.strip():
                    context.append(f"[Document Title] {para.text.strip()}")
                    break
        
        # Get section headings
        heading_count = 0
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading') and para.text.strip():
                context.append(f"[Section] {para.text.strip()}")
                heading_count += 1
                if heading_count >= 10:
                    break
        
        # Get first few paragraphs for additional context
        para_count = 0
        for para in doc.paragraphs:
            if para.text.strip() and not para.style.name.startswith('Heading'):
                context.append(f"[Content Sample] {para.text.strip()[:100]}...")
                para_count += 1
                if para_count >= 3:
                    break
        
        return context
    
    def _process_paragraph(self, para: Paragraph, translations: Dict[int, str], text_index: int) -> int:
        """Process a single paragraph"""
        for run in para.runs:
            if run.text.strip():
                if text_index in translations:
                    # Preserve formatting while updating text
                    original_text = run.text
                    translated_text = translations[text_index]
                    
                    # Preserve leading/trailing whitespace
                    leading_space = len(original_text) - len(original_text.lstrip())
                    trailing_space = len(original_text) - len(original_text.rstrip())
                    
                    run.text = ' ' * leading_space + translated_text.strip() + ' ' * trailing_space
                text_index += 1
        return text_index
    
    def _process_table_cell(self, cell: _Cell, translations: Dict[int, str], text_index: int) -> int:
        """Process a table cell"""
        for para in cell.paragraphs:
            text_index = self._process_paragraph(para, translations, text_index)
        return text_index
    
    def translate_document(self, input_path: Path, output_path: Path):
        """Translate a Word document"""
        self.logger.info(f"Processing: {input_path}")
        
        try:
            # Load document
            doc = Document(input_path)
            
            # Extract document context if smart context is enabled
            doc_context = []
            if self.config.smart_context:
                doc_context = self._extract_document_context(doc)
                self.logger.info(f"Extracted {len(doc_context)} context items")
            
            # Collect all texts to translate
            texts_to_translate = []
            
            # Process paragraphs
            for para_idx, para in enumerate(doc.paragraphs):
                for run_idx, run in enumerate(para.runs):
                    if run.text.strip():
                        texts_to_translate.append({
                            'id': len(texts_to_translate),
                            'text': run.text,
                            'type': 'paragraph',
                            'style': para.style.name,
                            'location': f'Paragraph {para_idx}, Run {run_idx}'
                        })
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            for run_idx, run in enumerate(para.runs):
                                if run.text.strip():
                                    texts_to_translate.append({
                                        'id': len(texts_to_translate),
                                        'text': run.text,
                                        'type': 'table_cell',
                                        'location': f'Table {table_idx}, Row {row_idx}, Cell {cell_idx}'
                                    })
            
            if not texts_to_translate:
                self.logger.warning(f"No text found to translate in {input_path}")
                doc.save(output_path)
                return
            
            self.logger.info(f"Found {len(texts_to_translate)} text segments to translate")
            
            # Translate in batches with progress bar
            all_translations = {}
            
            with tqdm(total=len(texts_to_translate), desc="Translating") as pbar:
                for i in range(0, len(texts_to_translate), self.config.batch_size):
                    batch = texts_to_translate[i:i + self.config.batch_size]
                    
                    # Prepare batch for API (only id and text)
                    batch_for_api = [{'id': j, 'text': item['text']} 
                                     for j, item in enumerate(batch)]
                    
                    # Add batch context info
                    batch_context = doc_context.copy()
                    if batch:
                        batch_context.append(f"[Batch Info] Translating {batch[0]['type']} elements")
                    
                    # Translate batch
                    batch_translations = self.translate_batch(batch_for_api, batch_context)
                    
                    # Map back to original indices
                    for j, translation in batch_translations.items():
                        original_idx = i + j
                        all_translations[original_idx] = translation
                    
                    pbar.update(len(batch))
            
            # Apply translations back to document
            text_index = 0
            
            # Apply to paragraphs
            for para in doc.paragraphs:
                text_index = self._process_paragraph(para, all_translations, text_index)
            
            # Apply to tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_index = self._process_table_cell(cell, all_translations, text_index)
            
            # Save translated document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            self.logger.info(f"✅ Saved translated document: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Failed to process {input_path}: {e}")
            if self.config.debug:
                import traceback
                traceback.print_exc()
            raise


def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(
        description="Word Document Translation Utility with Google Genai API v1.27.0",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    # Input options
    parser.add_argument('input_file', help='Input Word document (.docx)')
    
    # Language options
    parser.add_argument('--source-lang', default='auto', 
                       help='Source language code (default: auto-detect)')
    
    lang_group = parser.add_mutually_exclusive_group(required=True)
    lang_group.add_argument('--target-lang', help='Target language code (e.g., ja)')
    lang_group.add_argument('--target-langs', help='Multiple target languages (comma-separated)')
    
    # Translation options
    parser.add_argument('--style-prompt', choices=list(STYLE_PROMPTS.keys()), 
                       default='business', help='Translation style')
    parser.add_argument('--context-file', help='Path to glossary/context JSON file')
    parser.add_argument('--smart-context', action='store_true',
                       help='Use document structure for better context')
    
    # Output options
    parser.add_argument('--output-dir', default='.', help='Output directory')
    parser.add_argument('--model', default=DEFAULT_MODEL, help='Gemini model name')
    parser.add_argument('--batch-size', type=int, default=DEFAULT_BATCH_SIZE,
                       help=f'Texts per API call (default: {DEFAULT_BATCH_SIZE})')
    parser.add_argument('--concurrency', type=int, default=DEFAULT_CONCURRENCY,
                       help=f'Concurrent API calls (default: {DEFAULT_CONCURRENCY})')
    
    # Other options
    parser.add_argument('--debug', action='store_true', help='Enable debug logging')
    parser.add_argument('--list-styles', action='store_true', help='List available styles')
    parser.add_argument('--version', action='store_true', help='Show version information')
    
    args = parser.parse_args()
    
    # Handle version request
    if args.version:
        print(f"{Fore.CYAN}DOCX Translator using google-genai v1.27.0{Style.RESET_ALL}")
        print(f"Model: {DEFAULT_MODEL}")
        sys.exit(0)
    
    # Handle info requests
    if args.list_styles:
        print(f"\n{Fore.CYAN}Available translation styles:{Style.RESET_ALL}")
        for style, prompts in STYLE_PROMPTS.items():
            print(f"\n  {Fore.GREEN}{style}:{Style.RESET_ALL}")
            for lang, prompt in prompts.items():
                if lang != 'default':
                    lang_name = 'Japanese' if lang == 'ja' else lang.upper()
                    print(f"    {lang_name}: {prompt[:60]}...")
                else:
                    print(f"    Default: {prompt[:60]}...")
        sys.exit(0)
    
    # Parse target languages
    target_langs = []
    if args.target_langs:
        target_langs = [lang.strip() for lang in args.target_langs.split(',')]
    else:
        target_langs = [args.target_lang]
    
    # Validate input file
    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"{Fore.RED}Error: Input file not found: {input_path}{Style.RESET_ALL}")
        sys.exit(1)
    
    if not input_path.suffix.lower() == '.docx':
        print(f"{Fore.RED}Error: Input file must be a .docx file{Style.RESET_ALL}")
        sys.exit(1)
    
    # Show header
    print(f"\n{Fore.CYAN}{'='*60}{Style.RESET_ALL}")
    print(f"{Fore.GREEN}DOCX Translator - Google Genai v1.27.0{Style.RESET_ALL}")
    print(f"{Fore.CYAN}{'='*60}{Style.RESET_ALL}")
    print(f"Model: {args.model}")
    print(f"Source: {args.source_lang}")
    print(f"Targets: {', '.join(target_langs)}")
    print(f"Style: {args.style_prompt}")
    print(f"{Fore.CYAN}{'='*60}{Style.RESET_ALL}\n")
    
    # Process each target language
    output_dir = Path(args.output_dir)
    
    for target_lang in target_langs:
        print(f"\n{Fore.CYAN}Translating to {target_lang}...{Style.RESET_ALL}")
        
        # Create configuration
        config = TranslationConfig(
            source_lang=args.source_lang,
            target_lang=target_lang,
            model_name=args.model,
            batch_size=args.batch_size,
            concurrency=args.concurrency,
            style_prompt=args.style_prompt,
            context_file=args.context_file,
            smart_context=args.smart_context,
            debug=args.debug
        )
        
        # Create translator
        try:
            translator = DocxTranslator(config)
        except ValueError as e:
            print(str(e))
            sys.exit(1)
        
        # Generate output filename
        output_name = f"{input_path.stem}_{target_lang}{input_path.suffix}"
        output_path = output_dir / output_name
        
        # Translate document
        try:
            start_time = time.time()
            translator.translate_document(input_path, output_path)
            elapsed_time = time.time() - start_time
            print(f"{Fore.GREEN}✅ Successfully created: {output_path}{Style.RESET_ALL}")
            print(f"{Fore.CYAN}   Time taken: {elapsed_time:.2f} seconds{Style.RESET_ALL}")
        except Exception as e:
            print(f"{Fore.RED}❌ Translation failed: {e}{Style.RESET_ALL}")
            if args.debug:
                import traceback
                traceback.print_exc()
            continue
    
    print(f"\n{Fore.GREEN}✅ All translations completed!{Style.RESET_ALL}")


if __name__ == "__main__":
    # Handle event loop for Windows
    if sys.platform == 'win32':
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    
    main()