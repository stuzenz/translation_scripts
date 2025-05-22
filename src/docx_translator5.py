import argparse
import shutil
import json
import os
import re
from pathlib import Path
from docx import Document
import google.generativeai as genai
from colorama import Fore, Style, init

# Initialize colorama
init(autoreset=True)

def extract_json(response_text):
    """Extract JSON from markdown code blocks or raw text"""
    # Try to find JSON code blocks
    matches = re.findall(r'```(?:json)?\n(.*?)\n```', response_text, re.DOTALL)
    if matches:
        return matches[0]
    # If no code blocks, try to find first JSON structure
    match = re.search(r'{(.*?)}', response_text, re.DOTALL)
    if match:
        return f'{{{match.group(1)}}}'
    return response_text

def list_document_styles(input_file):
    """Lists all styles used in the document in alphabetical order."""
    input_path = Path(input_file)
    doc = Document(input_path)
    styles_in_doc = set()

    def collect_styles_from_element(element): # Renamed function for clarity
        if hasattr(element, 'style'):
            style_name = element.style.name
            if style_name:
                styles_in_doc.add(style_name)

    # Main document content - Iterate through paragraphs and tables directly
    for paragraph in doc.paragraphs:
        collect_styles_from_element(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    collect_styles_from_element(paragraph)

    # Headers and Footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    collect_styles_from_element(paragraph)
                for table in header.tables: # Added tables in headers
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                collect_styles_from_element(paragraph)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    collect_styles_from_element(paragraph)
                for table in footer.tables: # Added tables in footers
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                collect_styles_from_element(paragraph)


    sorted_styles = sorted(list(styles_in_doc), key=str.lower) # Alphabetical sort, case-insensitive
    print(f"{Fore.CYAN}🎉 Styles found in document '{input_file}':{Style.RESET_ALL}")
    for style in sorted_styles:
        print(f"\"{style}\"") # Print each style in quotes


def translate_file(input_file, source_lang, target_lang, model_name, batch_size, not_to_translate_styles=None):
    input_path = Path(input_file)
    # Construct output path with target language tag
    output_path = input_path.with_stem(f"{input_path.stem}_{target_lang}")
    shutil.copyfile(input_path, output_path)

    doc = Document(output_path)

    # Collect all text runs with indexes and styles
    all_runs = []

    def collect_runs(element):
        runs_with_style = []
        if hasattr(element, 'runs'):
            for run in element.runs:
                style_name = element.style.name if hasattr(element, 'style') else None # Paragraph Style
                runs_with_style.append({'run': run, 'style': style_name})
            return runs_with_style
        if hasattr(element, 'paragraphs'):
            for para in element.paragraphs:
                para_style_name = para.style.name if hasattr(para, 'style') else None # Paragraph Style
                for run in para.runs:
                    runs_with_style.append({'run': run, 'style': para_style_name})
            return runs_with_style
        if hasattr(element, 'rows'):
            for row in element.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_style_name = para.style.name if hasattr(para, 'style') else None # Paragraph Style
                        for run in para.runs:
                             runs_with_style.append({'run': run, 'style': para_style_name})
            return runs_with_style
        return []

    # Main document content
    for element in doc.paragraphs: # paragraphs for main document
        all_runs += collect_runs(element)

    # Process tables
    for table in doc.tables:
        all_runs += collect_runs(table)

    # Process headers/footers
    for section in doc.sections:
        for part in [section.header, section.footer,
                    section.first_page_header, section.first_page_footer,
                    section.even_page_header, section.even_page_footer]:
            if part is not None:
                for element in part.paragraphs: # paragraphs for headers/footers
                    all_runs += collect_runs(element)


    # Filter non-empty runs and apply style filter
    text_runs = []
    for i, run_data in enumerate(all_runs):
        run = run_data['run']
        style_name = run_data['style']
        if run.text.strip(): # Only process non-empty runs
            if not not_to_translate_styles or style_name not in not_to_translate_styles:
                text_runs.append((i, run)) # Keep original index, but now using filtered runs
            else:
                print(f"{Fore.YELLOW}⏭️ Skipping style '{style_name}': '{run.text[:50]}...'{Style.RESET_ALL}")

    # Initialize Gemini
    genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))
    model = genai.GenerativeModel(model_name)

    print(f"{Fore.CYAN}🚀 Processing {len(text_runs)} text elements in {len(text_runs)//batch_size + 1} batches{Style.RESET_ALL}")

    for batch_start in range(0, len(text_runs), batch_size):
        batch = text_runs[batch_start:batch_start + batch_size]
        batch_num = (batch_start // batch_size) + 1
        print(f"{Fore.YELLOW}🔧 Batch {batch_num} ({len(batch)} elements){Style.RESET_ALL}")

        batch_data = [{"id": idx, "text": run.text} for idx, run in batch]

        prompt = f"""
        Translate from {source_lang} to {target_lang}. Maintain formatting EXACTLY.
        Return ONLY VALID JSON using this format:
        {{
            "translations": [
                {{"id": <original_id>, "translation": "<translated_text>"}}
            ]
        }}
        DO NOT USE MARKDOWN. Ensure proper JSON escaping.
        The topic material for translation is a terms of reference document for an interim state architecture.
        Input: {json.dumps(batch_data, ensure_ascii=False)}
        """

        try:
            response = model.generate_content(prompt)
            print(f"{Fore.GREEN}📥 Raw Response:{Style.RESET_ALL} {response.text[:150]}...")

            # Clean and parse response
            cleaned_response = extract_json(response.text)
            print(f"{Fore.BLUE}🧹 Cleaned Response:{Style.RESET_ALL} {cleaned_response[:150]}...")

            try:
                result = json.loads(cleaned_response)
            except json.JSONDecodeError as e:
                print(f"{Fore.RED}❌ JSON Error:{Style.RESET_ALL} {str(e)}")
                print(f"{Fore.MAGENTA}🧬 Response Fragment:{Style.RESET_ALL} {cleaned_response[:500]}")
                continue

            if 'translations' not in result:
                print(f"{Fore.RED}⚠️ Missing 'translations' key{Style.RESET_ALL}")
                continue

            # Apply translations
            success = 0
            for item in result['translations']:
                try:
                    original_idx = item['id']
                    if original_idx >= len(all_runs): # Still check against the full all_runs length
                        print(f"{Fore.RED}🚨 Invalid ID {item['id']} in batch{Style.RESET_ALL}")
                        continue

                    # Find the correct run from the original all_runs list using the index
                    original_run = all_runs[original_idx]['run']
                    original_run.text = item.get('translation', original_run.text)
                    success += 1
                except Exception as e:
                    print(f"{Fore.RED}⚠️ Item error: {str(e)}{Style.RESET_ALL}")

            print(f"{Fore.GREEN}✅ Applied {success}/{len(batch)} translations{Style.RESET_ALL}")

        except Exception as e:
            print(f"{Fore.RED}💥 Batch Error:{Style.RESET_ALL} {str(e)}")
            print(f"{Fore.MAGENTA}📋 Failed Batch:{Style.RESET_ALL} {batch_data}")

    doc.save(output_path)
    print(f"{Fore.CYAN}🎉 Translation saved to: {output_path}{Style.RESET_ALL}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="DOCX Translator with Gemini")
    parser.add_argument("input_file", help="Input DOCX file path")
    parser.add_argument("--source-lang", default="auto", help="Source language code for translation")
    parser.add_argument("--target-lang", default=None, help="Target language code for translation") # Made target_lang optional if listing styles
    parser.add_argument("--model", default="gemini-2.0-flash", help="Gemini model name for translation")
    parser.add_argument("--batch-size", type=int, default=15,
                      help="Elements per batch for translation (15-20 recommended)")
    parser.add_argument("--not-to-translate-styles", default=None,
                      help="Comma-separated list of styles to skip translation for")
    parser.add_argument("--list-styles", action="store_true", help="List styles in the document and exit, do not translate") # New flag

    args = parser.parse_args()

    if args.list_styles:
        list_document_styles(args.input_file) # Call style listing function
    else: # Proceed with translation if --list-styles is not used
        if not args.target_lang: # target_lang is now required for translation only
            parser.error("--target-lang is required unless --list-styles is used.")

        not_to_translate_styles_list = None
        if args.not_to_translate_styles:
            not_to_translate_styles_list = [style.strip() for style in args.not_to_translate_styles.split(',')]

        translate_file(
            args.input_file,
            args.source_lang,
            args.target_lang,
            args.model,
            args.batch_size,
            not_to_translate_styles_list
        )
