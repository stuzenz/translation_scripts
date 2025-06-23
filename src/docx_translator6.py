import argparse
import shutil
import json
import os
import re
from pathlib import Path
from docx import Document
import google.generativeai as genai
from colorama import Fore, Style, init

# Initialize colorama
init(autoreset=True)

def extract_json(response_text):
    """Extract JSON from markdown code blocks or raw text"""
    # Try to find JSON code blocks
    matches = re.findall(r'```(?:json)?\n(.*?)\n```', response_text, re.DOTALL)
    if matches:
        return matches[0]
    # If no code blocks, try to find first JSON structure
    match = re.search(r'{(.*?)}', response_text, re.DOTALL)
    if match:
        return f'{{{match.group(1)}}}'
    return response_text

def list_document_styles(input_file):
    """Lists all styles used in the document in alphabetical order."""
    input_path = Path(input_file)
    doc = Document(input_path)
    styles_in_doc = set()

    def collect_styles_from_element(element):
        if hasattr(element, 'style'):
            style_name = element.style.name
            if style_name:
                styles_in_doc.add(style_name)

    # Main document content - Iterate through paragraphs and tables directly
    for paragraph in doc.paragraphs:
        collect_styles_from_element(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    collect_styles_from_element(paragraph)

    # Headers and Footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    collect_styles_from_element(paragraph)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                collect_styles_from_element(paragraph)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    collect_styles_from_element(paragraph)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                collect_styles_from_element(paragraph)

    sorted_styles = sorted(list(styles_in_doc), key=str.lower)
    print(f"{Fore.CYAN}🎉 Styles found in document '{input_file}':{Style.RESET_ALL}")
    for style in sorted_styles:
        print(f"\"{style}\"")


def translate_file_to_language(input_file, source_lang, target_lang, model, batch_size, not_to_translate_styles, output_dir):
    """Translate a single file to a single target language"""
    input_path = Path(input_file)
    
    # Determine output directory
    if output_dir:
        output_base_dir = Path(output_dir)
        output_base_dir.mkdir(parents=True, exist_ok=True)
    else:
        output_base_dir = input_path.parent
    
    # Construct output path with target language tag
    output_path = output_base_dir / f"{input_path.stem}_{target_lang}.docx"
    shutil.copyfile(input_path, output_path)

    doc = Document(output_path)

    # Collect all text runs with indexes and styles
    all_runs = []

    def collect_runs(element):
        runs_with_style = []
        if hasattr(element, 'runs'):
            for run in element.runs:
                style_name = element.style.name if hasattr(element, 'style') else None
                runs_with_style.append({'run': run, 'style': style_name})
            return runs_with_style
        if hasattr(element, 'paragraphs'):
            for para in element.paragraphs:
                para_style_name = para.style.name if hasattr(para, 'style') else None
                for run in para.runs:
                    runs_with_style.append({'run': run, 'style': para_style_name})
            return runs_with_style
        if hasattr(element, 'rows'):
            for row in element.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_style_name = para.style.name if hasattr(para, 'style') else None
                        for run in para.runs:
                             runs_with_style.append({'run': run, 'style': para_style_name})
            return runs_with_style
        return []

    # Main document content
    for element in doc.paragraphs:
        all_runs += collect_runs(element)

    # Process tables
    for table in doc.tables:
        all_runs += collect_runs(table)

    # Process headers/footers
    for section in doc.sections:
        for part in [section.header, section.footer,
                    section.first_page_header, section.first_page_footer,
                    section.even_page_header, section.even_page_footer]:
            if part is not None:
                for element in part.paragraphs:
                    all_runs += collect_runs(element)

    # Filter non-empty runs and apply style filter
    text_runs = []
    for i, run_data in enumerate(all_runs):
        run = run_data['run']
        style_name = run_data['style']
        if run.text.strip():
            if not not_to_translate_styles or style_name not in not_to_translate_styles:
                text_runs.append((i, run))
            else:
                print(f"{Fore.YELLOW}⏭️ Skipping style '{style_name}': '{run.text[:50]}...'{Style.RESET_ALL}")

    if not text_runs:
        print(f"{Fore.YELLOW}⚠️ No text to translate for {target_lang} in '{input_file}'{Style.RESET_ALL}")
        return output_path

    print(f"{Fore.CYAN}🚀 Processing {len(text_runs)} text elements in {len(text_runs)//batch_size + 1} batches for {target_lang}{Style.RESET_ALL}")

    for batch_start in range(0, len(text_runs), batch_size):
        batch = text_runs[batch_start:batch_start + batch_size]
        batch_num = (batch_start // batch_size) + 1
        print(f"{Fore.YELLOW}🔧 Batch {batch_num} ({len(batch)} elements) - {target_lang}{Style.RESET_ALL}")

        batch_data = [{"id": idx, "text": run.text} for idx, run in batch]

        prompt = f"""
        Translate from {source_lang} to {target_lang}. Maintain formatting EXACTLY.
        Return ONLY VALID JSON using this format:
        {{
            "translations": [
                {{"id": <original_id>, "translation": "<translated_text>"}}
            ]
        }}
        DO NOT USE MARKDOWN. Ensure proper JSON escaping.
        The topic material for translation is a terms of reference document for an interim state architecture.
        Input: {json.dumps(batch_data, ensure_ascii=False)}
        """

        try:
            response = model.generate_content(prompt)
            print(f"{Fore.GREEN}📥 Raw Response ({target_lang}):{Style.RESET_ALL} {response.text[:150]}...")

            # Clean and parse response
            cleaned_response = extract_json(response.text)
            print(f"{Fore.BLUE}🧹 Cleaned Response ({target_lang}):{Style.RESET_ALL} {cleaned_response[:150]}...")

            try:
                result = json.loads(cleaned_response)
            except json.JSONDecodeError as e:
                print(f"{Fore.RED}❌ JSON Error ({target_lang}):{Style.RESET_ALL} {str(e)}")
                print(f"{Fore.MAGENTA}🧬 Response Fragment:{Style.RESET_ALL} {cleaned_response[:500]}")
                continue

            if 'translations' not in result:
                print(f"{Fore.RED}⚠️ Missing 'translations' key ({target_lang}){Style.RESET_ALL}")
                continue

            # Apply translations
            success = 0
            for item in result['translations']:
                try:
                    original_idx = item['id']
                    if original_idx >= len(all_runs):
                        print(f"{Fore.RED}🚨 Invalid ID {item['id']} in batch ({target_lang}){Style.RESET_ALL}")
                        continue

                    original_run = all_runs[original_idx]['run']
                    original_run.text = item.get('translation', original_run.text)
                    success += 1
                except Exception as e:
                    print(f"{Fore.RED}⚠️ Item error ({target_lang}): {str(e)}{Style.RESET_ALL}")

            print(f"{Fore.GREEN}✅ Applied {success}/{len(batch)} translations ({target_lang}){Style.RESET_ALL}")

        except Exception as e:
            print(f"{Fore.RED}💥 Batch Error ({target_lang}):{Style.RESET_ALL} {str(e)}")

    doc.save(output_path)
    print(f"{Fore.CYAN}🎉 Translation for {target_lang} saved to: {output_path}{Style.RESET_ALL}")
    return output_path


def translate_file(input_file, source_lang, target_langs, model_name, batch_size, not_to_translate_styles=None, output_dir=None):
    """Translate a single file to multiple target languages"""
    input_path = Path(input_file)
    print(f"{Fore.MAGENTA}📄 Processing file: {input_path.name}{Style.RESET_ALL}")
    
    # Initialize Gemini model once
    genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))
    model = genai.GenerativeModel(model_name)
    
    output_files = []
    
    # Process each target language
    for target_lang in target_langs:
        print(f"{Fore.MAGENTA}🌍 Translating to {target_lang}...{Style.RESET_ALL}")
        try:
            output_file = translate_file_to_language(
                input_file, source_lang, target_lang, model, 
                batch_size, not_to_translate_styles, output_dir
            )
            output_files.append(output_file)
        except Exception as e:
            print(f"{Fore.RED}💥 Error translating to {target_lang}: {str(e)}{Style.RESET_ALL}")
    
    return output_files


def get_docx_files_from_directory(directory):
    """Get all DOCX files from a directory"""
    dir_path = Path(directory)
    if not dir_path.exists():
        raise FileNotFoundError(f"Directory '{directory}' does not exist")
    
    # Get all .docx files, excluding temporary files (starting with ~$)
    docx_files = [f for f in dir_path.glob("*.docx") if not f.name.startswith("~$")]
    
    if not docx_files:
        print(f"{Fore.YELLOW}⚠️ No DOCX files found in '{directory}'{Style.RESET_ALL}")
    else:
        print(f"{Fore.CYAN}📁 Found {len(docx_files)} DOCX files in '{directory}'{Style.RESET_ALL}")
        for file in docx_files:
            print(f"   📄 {file.name}")
    
    return docx_files


def main():
    parser = argparse.ArgumentParser(
        description="DOCX Translator with Gemini - Multi-language and batch processing support",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Translate single file to multiple languages
  python script.py document.docx --target-langs ja ko zh-cn
  
  # Process all DOCX files in a directory
  python script.py --input-files ./docs --target-langs ja ko --output-files ./translated
  
  # Using comma-separated format (also supported)
  python script.py document.docx --target-langs ja,ko,zh-cn
  
  # Using multiple --target-lang arguments
  python script.py document.docx --target-lang ja --target-lang ko
  
  # List styles in documents
  python script.py document.docx --list-styles
  python script.py --input-files ./docs --list-styles
  
  # Skip certain styles from translation
  python script.py document.docx --target-langs ja --not-to-translate-styles "Code,Heading 1"
        """
    )
    
    # Input options
    parser.add_argument("input_file", nargs='?', help="Single input DOCX file path")
    parser.add_argument("--input-files", help="Directory containing DOCX files to process")
    
    # Translation options
    parser.add_argument("--source-lang", default="auto", help="Source language code for translation")
    parser.add_argument("--target-lang", action="append", dest="target_langs_list", 
                      help="Target language code for translation (can be used multiple times)")
    parser.add_argument("--target-langs", nargs='+', help="Space-separated list of target language codes (e.g., 'ja ko zh-cn') or comma-separated string")
    parser.add_argument("--model", default="gemini-2.0-flash", help="Gemini model name for translation")
    parser.add_argument("--batch-size", type=int, default=15,
                      help="Elements per batch for translation (15-20 recommended)")
    parser.add_argument("--not-to-translate-styles", default=None,
                      help="Comma-separated list of styles to skip translation for")
    
    # Output options
    parser.add_argument("--output-files", help="Directory to save translated files (default: same as input)")
    
    # Utility options
    parser.add_argument("--list-styles", action="store_true", 
                      help="List styles in the document(s) and exit, do not translate")

    args = parser.parse_args()

    # Validate input arguments
    if not args.input_file and not args.input_files:
        parser.error("Either provide a single input file or use --input-files for directory processing")
    
    if args.input_file and args.input_files:
        parser.error("Cannot use both single input file and --input-files together")

    # Handle input files
    input_files = []
    if args.input_file:
        input_path = Path(args.input_file)
        if not input_path.exists():
            print(f"{Fore.RED}❌ Input file '{args.input_file}' does not exist{Style.RESET_ALL}")
            return 1
        input_files = [args.input_file]
    elif args.input_files:
        try:
            input_files = get_docx_files_from_directory(args.input_files)
            if not input_files:
                return 1
        except FileNotFoundError as e:
            print(f"{Fore.RED}❌ {str(e)}{Style.RESET_ALL}")
            return 1

    # Handle list styles option
    if args.list_styles:
        for input_file in input_files:
            print(f"\n{Fore.BLUE}📄 File: {Path(input_file).name}{Style.RESET_ALL}")
            try:
                list_document_styles(input_file)
            except Exception as e:
                print(f"{Fore.RED}❌ Error processing '{input_file}': {str(e)}{Style.RESET_ALL}")
        return 0

    # Handle target languages
    target_langs = []
    if args.target_langs_list:
        target_langs.extend(args.target_langs_list)
    if args.target_langs:
        # Handle both space-separated and comma-separated formats
        for lang_item in args.target_langs:
            if ',' in lang_item:
                # If comma-separated, split by comma
                target_langs.extend([lang.strip() for lang in lang_item.split(',')])
            else:
                # If space-separated (or single item), add as-is
                target_langs.append(lang_item.strip())
    
    if not target_langs:
        parser.error("At least one target language must be specified using --target-lang or --target-langs")

    # Remove duplicates while preserving order
    target_langs = list(dict.fromkeys(target_langs))
    print(f"{Fore.CYAN}🎯 Target languages: {', '.join(target_langs)}{Style.RESET_ALL}")

    # Handle not-to-translate styles
    not_to_translate_styles_list = None
    if args.not_to_translate_styles:
        not_to_translate_styles_list = [style.strip() for style in args.not_to_translate_styles.split(',')]
        print(f"{Fore.CYAN}🚫 Styles to skip: {', '.join(not_to_translate_styles_list)}{Style.RESET_ALL}")

    # Validate API key
    if not os.getenv('GOOGLE_API_KEY'):
        print(f"{Fore.RED}❌ GOOGLE_API_KEY environment variable is not set{Style.RESET_ALL}")
        return 1

    # Process files
    all_output_files = []
    total_files = len(input_files)
    
    for i, input_file in enumerate(input_files, 1):
        print(f"\n{Fore.BLUE}📊 Processing file {i}/{total_files}: {Path(input_file).name}{Style.RESET_ALL}")
        try:
            output_files = translate_file(
                input_file,
                args.source_lang,
                target_langs,
                args.model,
                args.batch_size,
                not_to_translate_styles_list,
                args.output_files
            )
            all_output_files.extend(output_files)
        except Exception as e:
            print(f"{Fore.RED}💥 Error processing '{input_file}': {str(e)}{Style.RESET_ALL}")

    # Summary
    print(f"\n{Fore.GREEN}🎉 Translation completed!{Style.RESET_ALL}")
    print(f"{Fore.CYAN}📊 Summary:{Style.RESET_ALL}")
    print(f"   📄 Input files processed: {total_files}")
    print(f"   🌍 Target languages: {len(target_langs)} ({', '.join(target_langs)})")
    print(f"   📁 Output files created: {len(all_output_files)}")
    
    if all_output_files:
        print(f"\n{Fore.CYAN}📁 Output files:{Style.RESET_ALL}")
        for output_file in all_output_files:
            print(f"   ✅ {output_file}")

    return 0


if __name__ == "__main__":
    exit(main())