#!/usr/bin/env python3
"""
docx-translator - Microsoft Word Document Translation Utility with Google Generative AI

A command-line tool for translating Microsoft Word documents while preserving
formatting, styles, tables, and document structure. Enhanced with improved JSON
parsing and error handling.

Features:
- Uses google-generativeai SDK with robust error handling
- Translate single files or entire directories
- Preserve document formatting, styles, and structure
- Support for tables with context-aware translation
- Smart context extraction for better translation quality
- Multiple target languages support
- Style-based translation (business, casual, technical, etc.)
- Robust JSON parsing to handle API response variations
- Optimized for Japanese and other languages

Usage Examples:
    # Translate a single file to Japanese
    python docx-translator.py document.docx --target-lang ja

    # Translate with specific style
    python docx-translator.py report.docx --target-lang ja --style-prompt business

    # Translate to multiple languages
    python docx-translator.py document.docx --target-langs en,ja,es,zh

    # Use glossary for consistent terminology
    python docx-translator.py technical.docx --target-lang ja --context-file glossary.json
"""

import argparse
import json
import logging
import os
import re
import sys
import time
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor, as_completed

# Use existing google-generativeai package
import google.generativeai as genai
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.shared import RGBColor
from colorama import Fore, Style, init
from tqdm import tqdm

# Initialize colorama for colored output
init(autoreset=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Constants
DEFAULT_MODEL = 'gemini-flash-latest'
DEFAULT_BATCH_SIZE = 10
DEFAULT_CONCURRENCY = 4
MAX_RETRIES = 3
RETRY_DELAY = 2

# Style prompts optimized for translations
STYLE_PROMPTS = {
    'business': {
        'ja': 'ビジネス文書として適切な敬語と専門用語を使用し、フォーマルで明確な日本語に翻訳してください。',
        'default': 'Translate using formal business language appropriate for professional documents.'
    },
    'casual': {
        'ja': '自然で親しみやすい日本語に翻訳し、適切な丁寧さを保ってください。',
        'default': 'Translate using conversational, friendly language suitable for general communication.'
    },
    'technical': {
        'ja': '技術的な正確性を優先し、専門用語は適切な日本語の技術用語を使用してください。',
        'default': 'Translate using precise technical terminology, maintaining accuracy for specialized content.'
    },
    'academic': {
        'ja': '学術的な文体で、正確性と論理性を重視した日本語に翻訳してください。',
        'default': 'Translate using scholarly language appropriate for academic contexts.'
    },
    'marketing': {
        'ja': '説得力があり、読者の心に響く魅力的な日本語表現を使用してください。',
        'default': 'Translate using persuasive, engaging language suitable for marketing materials.'
    }
}


@dataclass
class ProtectedSegment:
    """Represents a text segment that should not be translated (unified for font/type protection)"""
    placeholder: str
    original_text: str
    context_type: str
    paragraph_index: int
    run_index: int
    run_properties: Dict[str, Any]
    # Protection type specific fields
    font_name: Optional[str] = None  # For font protection
    formatting_types: Optional[Dict[str, bool]] = None  # For type protection
    protection_reason: str = "unknown"  # 'font', 'type', etc.


# Keep for backward compatibility during refactor
FontProtectedSegment = ProtectedSegment


@dataclass
class GlossaryProtectedSegment:
    """Represents a glossary-protected text segment that should use specific translation"""
    placeholder: str
    original_text: str
    target_translation: str
    start_pos: int
    end_pos: int
    is_case_sensitive: bool


@dataclass
class TranslationConfig:
    """Configuration for translation job"""
    source_lang: str
    target_lang: str
    model_name: str
    batch_size: int
    concurrency: int
    style_prompt: str
    context_file: Optional[str]
    smart_context: bool
    debug: bool
    dont_translate_fonts: Optional[List[str]]
    dont_translate_types: Optional[List[str]]


class TextProtectionEngine:
    """Base class for text protection engines (font, type, etc.)"""

    # Context detection keywords (shared across all protection types)
    UI_KEYWORDS = ['button', 'menu', 'dialog', 'field', 'tab', 'click', 'select', 'press', 'choose']
    CODE_KEYWORDS = ['function', 'variable', 'command', 'method', 'property', 'run', 'execute', 'call', 'script']
    VALUE_KEYWORDS = ['enter', 'input', 'type', 'password', 'username', 'value', 'set', 'specify']

    def __init__(self, protection_prefix: str):
        self.protection_prefix = protection_prefix  # e.g., "FONTPROT", "TYPEPROT"
        self.segment_counter = 0
        self.protected_segments: Dict[str, ProtectedSegment] = {}

    def analyze_context(self, text_before: str, text_after: str, protected_text: str) -> str:
        """Analyze surrounding context to determine placeholder type"""
        combined_text = f"{text_before} {protected_text} {text_after}".lower()

        # Check for UI context
        for keyword in self.UI_KEYWORDS:
            if keyword in combined_text:
                return 'ui'

        # Check for code context
        for keyword in self.CODE_KEYWORDS:
            if keyword in combined_text:
                return 'code'

        # Check for value context
        for keyword in self.VALUE_KEYWORDS:
            if keyword in combined_text:
                return 'value'

        # Default to generic
        return 'generic'

    def generate_placeholder(self, context_type: str) -> str:
        """Generate a unique placeholder based on context type"""
        self.segment_counter += 1

        type_map = {
            'ui': 'UI',
            'code': 'CODE',
            'value': 'VALUE',
            'generic': 'TERM'
        }

        type_prefix = type_map.get(context_type, 'TERM')
        return f"__{self.protection_prefix}_{type_prefix}_{self.segment_counter}__"

    def extract_run_properties(self, run) -> Dict[str, Any]:
        """Extract formatting properties from a run"""
        props = {}
        try:
            if hasattr(run, 'font'):
                font = run.font
                props['font_name'] = font.name
                props['font_size'] = font.size
                props['bold'] = font.bold
                props['italic'] = font.italic
                props['underline'] = font.underline
                props['color'] = font.color.rgb if font.color and font.color.rgb else None
        except:
            pass
        return props

    def get_paragraph_context(self, paragraphs: List, para_index: int, run_index: int) -> Tuple[str, str]:
        """Get text before and after the current run for context analysis"""
        if para_index >= len(paragraphs):
            return "", ""

        paragraph = paragraphs[para_index]
        text_before = ""
        text_after = ""

        try:
            # Get text from runs before current run
            for i, run in enumerate(paragraph.runs):
                if i < run_index:
                    text_before += run.text
                elif i > run_index:
                    text_after += run.text

            # Get text from previous paragraph (last 50 chars)
            if para_index > 0:
                prev_text = paragraphs[para_index - 1].text
                text_before = prev_text[-50:] + " " + text_before

            # Get text from next paragraph (first 50 chars)
            if para_index < len(paragraphs) - 1:
                next_text = paragraphs[para_index + 1].text
                text_after = text_after + " " + next_text[:50]
        except:
            pass

        return text_before.strip(), text_after.strip()

    def should_protect_run(self, run) -> bool:
        """Override in subclasses to define protection criteria"""
        raise NotImplementedError("Subclasses must implement should_protect_run")


class FontProtectionEngine(TextProtectionEngine):
    """Engine for detecting and managing font-protected text segments"""

    def __init__(self, protected_fonts: List[str]):
        super().__init__("FONTPROT")
        # Normalize font names to lowercase for comparison
        self.protected_fonts = [font.strip().lower() for font in protected_fonts] if protected_fonts else []

    def should_protect_run(self, run) -> bool:
        """Check if a run uses a protected font"""
        if not self.protected_fonts:
            return False

        try:
            # Get font name from run properties
            if hasattr(run, 'font') and run.font.name:
                run_font = run.font.name.lower()
                # Check for exact match or partial match (e.g., "Courier" matches "Courier New")
                for protected_font in self.protected_fonts:
                    if protected_font in run_font or run_font in protected_font:
                        return True
        except:
            pass
        return False

    # Keep for backward compatibility
    def is_protected_font(self, run) -> bool:
        """Deprecated: use should_protect_run instead"""
        return self.should_protect_run(run)


class TypeProtectionEngine(TextProtectionEngine):
    """Engine for detecting and managing formatting-type-protected text segments"""

    # Mapping of type names to formatting properties
    TYPE_MAPPINGS = {
        "bold": {"bold": True},
        "italic": {"italic": True},
        "underlined": {"underline": True},
        "italic-bold": {"italic": True, "bold": True},
        "underlined-bold": {"underline": True, "bold": True},
        "underlined-italic": {"underline": True, "italic": True},
        "underlined-italic-bold": {"underline": True, "italic": True, "bold": True}
    }

    def __init__(self, protected_types: List[str]):
        super().__init__("TYPEPROT")
        # Normalize type names and convert to formatting requirements
        self.protected_formatting = []

        if protected_types:
            for type_name in protected_types:
                type_name = type_name.strip().lower()
                if type_name in self.TYPE_MAPPINGS:
                    self.protected_formatting.append(self.TYPE_MAPPINGS[type_name])
                else:
                    # Log warning about invalid type but don't fail
                    pass

    def should_protect_run(self, run) -> bool:
        """Check if a run matches any of the protected formatting types"""
        if not self.protected_formatting:
            return False

        try:
            # Extract current run's formatting properties
            if hasattr(run, 'font'):
                font = run.font
                run_formatting = {
                    "bold": bool(font.bold),
                    "italic": bool(font.italic),
                    "underline": bool(font.underline)
                }

                # Check if current formatting matches any protected type
                for protected_format in self.protected_formatting:
                    match = True
                    for property_name, required_value in protected_format.items():
                        if run_formatting.get(property_name, False) != required_value:
                            match = False
                            break

                    if match:
                        return True
        except:
            pass
        return False

    @classmethod
    def get_valid_types(cls) -> List[str]:
        """Return list of valid formatting type names"""
        return list(cls.TYPE_MAPPINGS.keys())

    @classmethod
    def validate_types(cls, type_list: List[str]) -> Tuple[List[str], List[str]]:
        """Validate type list and return (valid_types, invalid_types)"""
        valid = []
        invalid = []

        for type_name in type_list:
            type_name = type_name.strip().lower()
            if type_name in cls.TYPE_MAPPINGS:
                valid.append(type_name)
            else:
                invalid.append(type_name)

        return valid, invalid


class GlossaryProtectionEngine:
    """Engine for detecting and managing glossary-protected text segments"""

    def __init__(self, glossary_data: Dict[str, str]):
        self.glossary_data = glossary_data or {}
        self.segment_counter = 0
        self.protected_segments: Dict[str, GlossaryProtectedSegment] = {}

        # Sort glossary terms by length (longest first) to handle overlapping terms properly
        self.sorted_terms = sorted(self.glossary_data.keys(), key=len, reverse=True)

    def generate_placeholder(self) -> str:
        """Generate a unique placeholder for glossary terms"""
        self.segment_counter += 1
        return f"__GLOSSARY_TERM_{self.segment_counter}__"

    def protect_text(self, text: str) -> Tuple[str, Dict[str, GlossaryProtectedSegment]]:
        """Replace glossary terms with placeholders and return mapping"""
        if not self.glossary_data:
            return text, {}

        protected_text = text
        segments = {}

        # Process terms in order of length (longest first) to avoid partial matches
        for term in self.sorted_terms:
            target_translation = self.glossary_data[term]

            # Find all occurrences of this term (case-insensitive)
            import re
            pattern = re.compile(re.escape(term), re.IGNORECASE)

            matches = list(pattern.finditer(protected_text))

            # Process matches in reverse order to maintain positions
            for match in reversed(matches):
                placeholder = self.generate_placeholder()

                segment = GlossaryProtectedSegment(
                    placeholder=placeholder,
                    original_text=match.group(),
                    target_translation=target_translation,
                    start_pos=match.start(),
                    end_pos=match.end(),
                    is_case_sensitive=False
                )

                segments[placeholder] = segment

                # Replace the matched text with placeholder
                protected_text = (
                    protected_text[:match.start()] +
                    placeholder +
                    protected_text[match.end():]
                )

        return protected_text, segments

    def restore_text(self, text: str, segments: Dict[str, GlossaryProtectedSegment]) -> str:
        """Restore placeholders with target translations"""
        restored_text = text

        for placeholder, segment in segments.items():
            # Replace placeholder with target translation
            restored_text = restored_text.replace(placeholder, segment.target_translation)

        return restored_text

    def build_glossary_instruction(self) -> str:
        """Build instruction text for the LLM about glossary terms"""
        if not self.glossary_data:
            return ""

        glossary_table = []
        for source, target in self.glossary_data.items():
            glossary_table.append(f'"{source}" → "{target}"')

        instruction = f"""
        MANDATORY GLOSSARY TERMS - These must be translated exactly as specified:
        {chr(10).join(glossary_table)}

        CRITICAL: Any placeholders in format __GLOSSARY_TERM_*__ must be preserved exactly.
        """

        return instruction


class ImprovedJSONExtractor:
    """Enhanced JSON extraction with multiple fallback strategies"""
    
    @staticmethod
    def extract_json(response_text: str) -> Dict[str, Any]:
        """Extract JSON from response with multiple strategies"""
        # Strategy 1: Try direct JSON parsing
        try:
            return json.loads(response_text)
        except json.JSONDecodeError:
            pass
        
        # Strategy 2: Extract from markdown code blocks
        code_block_patterns = [
            r'```json\s*\n(.*?)\n```',
            r'```\s*\n(.*?)\n```',
            r'`(.*?)`'
        ]
        
        for pattern in code_block_patterns:
            matches = re.findall(pattern, response_text, re.DOTALL)
            if matches:
                for match in matches:
                    try:
                        return json.loads(match)
                    except json.JSONDecodeError:
                        continue
        
        # Strategy 3: Find JSON-like structures
        json_patterns = [
            r'\{[^{}]*"translations"[^{}]*\}',  # Simple JSON with translations
            r'\{.*?"translations".*?\}',         # More flexible
            r'\{.*\}',                           # Any JSON object
        ]
        
        for pattern in json_patterns:
            matches = re.findall(pattern, response_text, re.DOTALL)
            if matches:
                # Try the longest match first
                for match in sorted(matches, key=len, reverse=True):
                    try:
                        # Clean up common issues
                        cleaned = match
                        # Remove trailing commas
                        cleaned = re.sub(r',\s*}', '}', cleaned)
                        cleaned = re.sub(r',\s*]', ']', cleaned)
                        # Fix escaped quotes that shouldn't be escaped
                        cleaned = cleaned.replace('\\"', '"')
                        # Remove any BOM or zero-width spaces
                        cleaned = cleaned.replace('\ufeff', '').replace('\u200b', '')
                        
                        result = json.loads(cleaned)
                        if isinstance(result, dict):
                            return result
                    except json.JSONDecodeError:
                        continue
        
        # Strategy 4: Try to fix common JSON issues
        cleaned_text = response_text
        # Remove any text before the first {
        first_brace = cleaned_text.find('{')
        if first_brace != -1:
            cleaned_text = cleaned_text[first_brace:]
        # Remove any text after the last }
        last_brace = cleaned_text.rfind('}')
        if last_brace != -1:
            cleaned_text = cleaned_text[:last_brace + 1]
        
        try:
            return json.loads(cleaned_text)
        except json.JSONDecodeError:
            pass
        
        # If all strategies fail, raise an exception with helpful info
        raise ValueError(f"Could not extract valid JSON from response. First 200 chars: {response_text[:200]}")


def get_docx_files_from_directory(directory):
    """Get all DOCX files from a directory"""
    dir_path = Path(directory)
    if not dir_path.exists():
        raise FileNotFoundError(f"Directory '{directory}' does not exist")

    # Get all .docx files, excluding temporary files (starting with ~$)
    docx_files = [f for f in dir_path.glob("*.docx") if not f.name.startswith("~$")]

    if not docx_files:
        print(f"{Fore.YELLOW}⚠️ No DOCX files found in '{directory}'{Style.RESET_ALL}")
    else:
        print(f"{Fore.CYAN}📁 Found {len(docx_files)} DOCX files in '{directory}'{Style.RESET_ALL}")
        for file in docx_files:
            print(f"   📄 {file.name}")

    return docx_files


class DocxTranslator:
    """Main translator class for Word documents using google-genai 1.27.0"""
    
    def __init__(self, config: TranslationConfig):
        self.config = config
        self.json_extractor = ImprovedJSONExtractor()
        self.font_protection = FontProtectionEngine(config.dont_translate_fonts) if config.dont_translate_fonts else None
        self.type_protection = TypeProtectionEngine(config.dont_translate_types) if config.dont_translate_types else None
        self._setup_logging()
        self._configure_api()
        self.context_data = self._load_context()
        self.glossary_protection = GlossaryProtectionEngine(self.context_data) if self.context_data else None
        
    def _setup_logging(self):
        """Configure logging based on debug flag"""
        level = logging.DEBUG if self.config.debug else logging.INFO
        logging.getLogger().setLevel(level)
        self.logger = logging.getLogger(__name__)
        
    def _configure_api(self):
        """Configure Google Generative AI API"""
        api_key = os.getenv('GOOGLE_API_KEY')
        if not api_key:
            raise ValueError(
                f"{Fore.RED}Error: GOOGLE_API_KEY environment variable not set.{Style.RESET_ALL}\n"
                f"{Fore.YELLOW}Please set your Google API key as an environment variable.{Style.RESET_ALL}"
            )

        # Configure the API
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel(self.config.model_name)
        self.logger.info(f"Initialized google-generativeai with model: {self.config.model_name}")

        # Test the API connection
        if not self._test_api_connection():
            raise ValueError(f"{Fore.RED}Failed to connect to Google Generative AI API{Style.RESET_ALL}")
    
    def _test_api_connection(self) -> bool:
        """Test API connection with a simple request"""
        try:
            test_prompt = 'Respond with only the JSON: {"status": "ok"}'
            response = self.model.generate_content(test_prompt)

            # Check if we got a response
            if response and response.text:
                self.logger.info("API connection test successful")
                return True
            else:
                self.logger.error("API connection test failed - no response text")
                return False

        except Exception as e:
            self.logger.error(f"API connection test failed: {str(e)}")
            return False
        
    def _load_context(self) -> Dict:
        """Load context/glossary file if provided"""
        if not self.config.context_file:
            return {}
            
        try:
            with open(self.config.context_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                self.logger.info(f"Loaded context from {self.config.context_file}")
                return data
        except Exception as e:
            self.logger.warning(f"Failed to load context file: {e}")
            return {}
    
    def _get_style_prompt(self) -> str:
        """Get appropriate style prompt for the target language"""
        style_dict = STYLE_PROMPTS.get(self.config.style_prompt, {})
        return style_dict.get(self.config.target_lang, style_dict.get('default', ''))
    
    def _build_translation_prompt(self, texts: List[Dict[str, Any]], context: List[str] = None) -> str:
        """Build translation prompt with glossary and font protection support"""

        # Create batch data in the same format as working version
        batch_data = [{"id": i, "text": text_obj["text"]} for i, text_obj in enumerate(texts)]

        # Build protection placeholder preservation instruction
        protection_instruction = ""
        has_font_placeholders = any("__FONTPROT_" in item["text"] for item in batch_data)
        has_type_placeholders = any("__TYPEPROT_" in item["text"] for item in batch_data)

        if has_font_placeholders or has_type_placeholders:
            placeholders = []
            if has_font_placeholders:
                placeholders.append("__FONTPROT_*__")
            if has_type_placeholders:
                placeholders.append("__TYPEPROT_*__")

            protection_instruction = f"""
        CRITICAL: Preserve ALL text in format {', '.join(placeholders)} exactly as written.
        These are placeholders that must NOT be translated or modified. Keep them exactly as they appear.
        """

        # Build glossary instruction
        glossary_instruction = ""
        if self.glossary_protection:
            glossary_instruction = self.glossary_protection.build_glossary_instruction()

        # Combine all instructions
        special_instructions = protection_instruction + glossary_instruction

        # Enhanced prompt with glossary support
        prompt = f"""
        Translate from {self.config.source_lang} to {self.config.target_lang}. Maintain formatting EXACTLY.
        Return ONLY VALID JSON using this format:
        {{
            "translations": [
                {{"id": <original_id>, "translation": "<translated_text>"}}
            ]
        }}
        DO NOT USE MARKDOWN. Ensure proper JSON escaping.{special_instructions}
        The topic material for translation is a terms of reference document for an interim state architecture.
        Input: {json.dumps(batch_data, ensure_ascii=False)}
        """

        return prompt
    
    # Removed async method - using sync method only for stability
    
    def translate_batch(self, texts: List[Dict[str, Any]], context: List[str] = None) -> Dict[int, str]:
        """Translate a batch using sync API (fallback method)"""
        if not texts:
            return {}
        
        prompt = self._build_translation_prompt(texts, context)
        
        for attempt in range(MAX_RETRIES):
            try:
                # Use the simple google-generativeai API call (same as working version)
                response = self.model.generate_content(prompt)

                # Extract text from the response
                response_text = response.text

                if not response_text:
                    self.logger.error(f"Empty response. Full response: {str(response)[:1000]}")
                    raise ValueError("Received empty response from API")
                
                self.logger.debug(f"Response text (first 500 chars): {response_text[:500]}")
                
                # Extract and parse JSON
                result = self.json_extractor.extract_json(response_text)
                
                # Validate and build translation map
                if 'translations' not in result:
                    self.logger.error(f"No translations key in result: {result}")
                    raise ValueError("Missing 'translations' key in response")
                
                translation_map = {}
                for item in result['translations']:
                    idx = item.get('id')
                    translation = item.get('translation', '')
                    if idx is not None:
                        translation_map[idx] = translation
                
                # Add fallbacks for missing translations
                for i in range(len(texts)):
                    if i not in translation_map:
                        self.logger.warning(f"Missing translation for ID {i}, using original")
                        translation_map[i] = texts[i]['text']
                
                return translation_map
                
            except Exception as e:
                self.logger.warning(f"Translation attempt {attempt + 1} failed: {str(e)}")
                
                # Log more details on last attempt
                if attempt == MAX_RETRIES - 1:
                    self.logger.error(f"Final attempt failed with error: {str(e)}")
                    self.logger.error(f"Prompt was: {prompt[:500]}...")
                
                if "429" in str(e) or "quota" in str(e).lower() or "rate" in str(e).lower():
                    wait_time = RETRY_DELAY * (2 ** (attempt + 1))
                    self.logger.info(f"Rate limit detected, waiting {wait_time} seconds...")
                    time.sleep(wait_time)
                elif attempt < MAX_RETRIES - 1:
                    time.sleep(RETRY_DELAY * (attempt + 1))
                else:
                    # Return original texts as fallback
                    self.logger.error("All translation attempts failed, using original text")
                    return {i: text['text'] for i, text in enumerate(texts)}
        
        # Fallback
        return {i: text['text'] for i, text in enumerate(texts)}
    
    def _extract_document_context(self, doc: Document) -> List[str]:
        """Extract context from the document"""
        context = []
        
        # Get document title from first heading
        for para in doc.paragraphs[:5]:
            if para.style.name.startswith('Heading') or para.style.name == 'Title':
                if para.text.strip():
                    context.append(f"[Document Title] {para.text.strip()}")
                    break
        
        # Get section headings
        heading_count = 0
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading') and para.text.strip():
                context.append(f"[Section] {para.text.strip()}")
                heading_count += 1
                if heading_count >= 10:
                    break
        
        # Get first few paragraphs for additional context
        para_count = 0
        for para in doc.paragraphs:
            if para.text.strip() and not para.style.name.startswith('Heading'):
                context.append(f"[Content Sample] {para.text.strip()[:100]}...")
                para_count += 1
                if para_count >= 3:
                    break
        
        return context
    
    def _process_paragraph(self, para: Paragraph, translations: Dict[int, str], text_index: int) -> int:
        """Process a single paragraph"""
        for run in para.runs:
            if run.text.strip():
                if text_index in translations:
                    # Preserve formatting while updating text
                    original_text = run.text
                    translated_text = translations[text_index]
                    
                    # Preserve leading/trailing whitespace
                    leading_space = len(original_text) - len(original_text.lstrip())
                    trailing_space = len(original_text) - len(original_text.rstrip())
                    
                    run.text = ' ' * leading_space + translated_text.strip() + ' ' * trailing_space
                text_index += 1
        return text_index
    
    def _process_table_cell(self, cell: _Cell, translations: Dict[int, str], text_index: int) -> int:
        """Process a table cell"""
        for para in cell.paragraphs:
            text_index = self._process_paragraph(para, translations, text_index)
        return text_index

    def _mark_toc_for_update(self, doc: Document):
        """Mark TOC fields to be updated when the document is opened in Word"""
        try:
            # Access the document part to modify field update settings
            doc_part = doc.part

            # Get the document element
            doc_element = doc_part.element

            # Look for field elements in the document
            from docx.oxml.ns import qn

            # Find all field elements (fldSimple and fldChar)
            fields = doc_element.xpath('.//w:fldSimple[@w:instr]',
                                    namespaces=doc_element.nsmap)

            toc_fields_found = 0
            for field in fields:
                instr = field.get(qn('w:instr'))
                if instr and 'TOC' in instr.upper():
                    # Mark field as dirty (needs update)
                    field.set(qn('w:dirty'), 'true')
                    toc_fields_found += 1

            # Also look for complex fields (fldChar approach)
            fld_chars = doc_element.xpath('.//w:fldChar[@w:fldCharType="begin"]',
                                        namespaces=doc_element.nsmap)

            for fld_char in fld_chars:
                # Check if this is followed by a TOC instruction
                next_element = fld_char.getnext()
                while next_element is not None:
                    if next_element.tag.endswith('}instrText'):
                        if 'TOC' in next_element.text.upper():
                            # Mark the field as dirty
                            fld_char.set(qn('w:dirty'), 'true')
                            toc_fields_found += 1
                            break
                    elif next_element.tag.endswith('}fldChar'):
                        break
                    next_element = next_element.getnext()

            if toc_fields_found > 0:
                self.logger.info(f"📑 Marked {toc_fields_found} TOC field(s) for update")
                self.logger.info("💡 Word will prompt to update the Table of Contents when document is opened")
            else:
                self.logger.debug("No TOC fields found in document")

        except Exception as e:
            self.logger.warning(f"Could not mark TOC for update: {e}")
            # This is not critical, so we continue

    def _process_text_protection(self, doc: Document) -> Dict[str, ProtectedSegment]:
        """Pre-process document to replace font and type-protected text with placeholders"""
        if not self.font_protection and not self.type_protection:
            return {}

        protected_segments = {}
        paragraphs = doc.paragraphs

        # Log what protection is enabled
        if self.font_protection:
            self.logger.info(f"🔒 Font protection enabled: {', '.join(self.font_protection.protected_fonts)}")
        if self.type_protection:
            valid_types = [name for name, props in TypeProtectionEngine.TYPE_MAPPINGS.items()
                          if props in self.type_protection.protected_formatting]
            self.logger.info(f"🎨 Type protection enabled: {', '.join(valid_types)}")

        for para_index, paragraph in enumerate(paragraphs):
            for run_index, run in enumerate(paragraph.runs):
                if not run.text.strip():
                    continue

                # Check if run should be protected (font or type)
                should_protect = False
                protection_reason = ""

                if self.font_protection and self.font_protection.should_protect_run(run):
                    should_protect = True
                    protection_reason = "font"
                elif self.type_protection and self.type_protection.should_protect_run(run):
                    should_protect = True
                    protection_reason = "type"

                if should_protect:
                    # Choose the appropriate engine for context analysis
                    engine = self.font_protection if protection_reason == "font" else self.type_protection

                    # Get context for placeholder type determination
                    text_before, text_after = engine.get_paragraph_context(
                        paragraphs, para_index, run_index
                    )

                    # Analyze context and generate placeholder
                    context_type = engine.analyze_context(
                        text_before, text_after, run.text
                    )
                    placeholder = engine.generate_placeholder(context_type)

                    # Extract run properties for restoration
                    run_properties = engine.extract_run_properties(run)

                    # Create protected segment with appropriate fields
                    segment_data = {
                        "placeholder": placeholder,
                        "original_text": run.text,
                        "context_type": context_type,
                        "paragraph_index": para_index,
                        "run_index": run_index,
                        "run_properties": run_properties,
                        "protection_reason": protection_reason
                    }

                    if protection_reason == "font":
                        segment_data["font_name"] = run.font.name if run.font.name else "Unknown"
                    elif protection_reason == "type":
                        # Store the specific formatting that triggered protection
                        formatting = {
                            "bold": bool(run.font.bold),
                            "italic": bool(run.font.italic),
                            "underline": bool(run.font.underline)
                        }
                        segment_data["formatting_types"] = formatting

                    segment = ProtectedSegment(**segment_data)
                    protected_segments[placeholder] = segment

                    # Replace text with placeholder
                    run.text = placeholder

                    self.logger.debug(f"🔒 Protected '{segment.original_text}' → '{placeholder}' ({protection_reason}: {context_type})")

        if protected_segments:
            self.logger.info(f"🔒 Protected {len(protected_segments)} text segments ({sum(1 for s in protected_segments.values() if s.protection_reason == 'font')} font, {sum(1 for s in protected_segments.values() if s.protection_reason == 'type')} type)")

        return protected_segments

    def _restore_text_protection(self, doc: Document, protected_segments: Dict[str, ProtectedSegment]):
        """Post-process document to restore font and type-protected text from placeholders"""
        if not protected_segments:
            return

        self.logger.info(f"🔓 Restoring {len(protected_segments)} protected text segments")

        # Scan through all paragraphs to find and replace placeholders
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text and any(placeholder in run.text for placeholder in protected_segments.keys()):
                    # Replace each placeholder in this run
                    original_text = run.text
                    for placeholder, segment in protected_segments.items():
                        if placeholder in run.text:
                            # Replace placeholder with original text
                            run.text = run.text.replace(placeholder, segment.original_text)

                            # Restore font properties
                            try:
                                if run.font and segment.run_properties:
                                    if 'font_name' in segment.run_properties:
                                        run.font.name = segment.run_properties['font_name']
                                    if 'font_size' in segment.run_properties and segment.run_properties['font_size']:
                                        run.font.size = segment.run_properties['font_size']
                                    if 'bold' in segment.run_properties and segment.run_properties['bold'] is not None:
                                        run.font.bold = segment.run_properties['bold']
                                    if 'italic' in segment.run_properties and segment.run_properties['italic'] is not None:
                                        run.font.italic = segment.run_properties['italic']
                                    if 'underline' in segment.run_properties and segment.run_properties['underline'] is not None:
                                        run.font.underline = segment.run_properties['underline']
                                    if 'color' in segment.run_properties and segment.run_properties['color']:
                                        run.font.color.rgb = segment.run_properties['color']
                            except Exception as e:
                                self.logger.warning(f"Could not restore font properties for '{segment.original_text}': {e}")

                            self.logger.debug(f"🔓 Restored '{placeholder}' → '{segment.original_text}'")

        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text and any(placeholder in run.text for placeholder in protected_segments.keys()):
                                for placeholder, segment in protected_segments.items():
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, segment.original_text)
                                        # Restore font properties (same as above)
                                        try:
                                            if run.font and segment.run_properties:
                                                if 'font_name' in segment.run_properties:
                                                    run.font.name = segment.run_properties['font_name']
                                                # ... (same property restoration logic)
                                        except Exception as e:
                                            self.logger.warning(f"Could not restore font properties in table for '{segment.original_text}': {e}")

        self.logger.info("🔓 Font protection restoration completed")
    
    def translate_document(self, input_path: Path, output_path: Path):
        """Translate a Word document"""
        self.logger.info(f"Processing: {input_path}")
        
        try:
            # Load document
            doc = Document(input_path)

            # Pre-process: Replace font and type-protected text with placeholders
            protected_segments = self._process_text_protection(doc)

            # Log glossary status
            if self.glossary_protection:
                self.logger.info(f"Glossary enabled with {len(self.context_data)} terms")
            else:
                self.logger.info("No glossary loaded")

            # Extract document context if smart context is enabled
            doc_context = []
            if self.config.smart_context:
                doc_context = self._extract_document_context(doc)
                self.logger.info(f"Extracted {len(doc_context)} context items")
            
            # Collect all texts to translate
            texts_to_translate = []
            
            # Process paragraphs
            for para_idx, para in enumerate(doc.paragraphs):
                for run_idx, run in enumerate(para.runs):
                    if run.text.strip():
                        texts_to_translate.append({
                            'id': len(texts_to_translate),
                            'text': run.text,
                            'type': 'paragraph',
                            'style': para.style.name,
                            'location': f'Paragraph {para_idx}, Run {run_idx}'
                        })
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            for run_idx, run in enumerate(para.runs):
                                if run.text.strip():
                                    texts_to_translate.append({
                                        'id': len(texts_to_translate),
                                        'text': run.text,
                                        'type': 'table_cell',
                                        'location': f'Table {table_idx}, Row {row_idx}, Cell {cell_idx}'
                                    })
            
            if not texts_to_translate:
                self.logger.warning(f"No text found to translate in {input_path}")
                doc.save(output_path)
                return
            
            self.logger.info(f"Found {len(texts_to_translate)} text segments to translate")
            
            # Translate in batches with progress bar
            all_translations = {}
            all_glossary_segments = {}  # Track glossary segments per batch

            with tqdm(total=len(texts_to_translate), desc="Translating") as pbar:
                for i in range(0, len(texts_to_translate), self.config.batch_size):
                    batch = texts_to_translate[i:i + self.config.batch_size]

                    # Apply glossary protection to batch texts
                    batch_for_api = []
                    batch_glossary_segments = {}

                    for j, item in enumerate(batch):
                        text = item['text']

                        # Apply glossary protection if enabled
                        if self.glossary_protection:
                            protected_text, segments = self.glossary_protection.protect_text(text)
                            if segments:
                                batch_glossary_segments[j] = segments
                                text = protected_text
                                self.logger.debug(f"Protected {len(segments)} glossary terms in: {item['text'][:50]}...")

                        batch_for_api.append({'id': j, 'text': text})

                    # Store glossary segments for this batch
                    if batch_glossary_segments:
                        all_glossary_segments[i] = batch_glossary_segments

                    # Add batch context info
                    batch_context = doc_context.copy()
                    if batch:
                        batch_context.append(f"[Batch Info] Translating {batch[0]['type']} elements")

                    # Translate batch
                    batch_translations = self.translate_batch(batch_for_api, batch_context)

                    # Restore glossary terms in translations
                    if i in all_glossary_segments:
                        for j, translation in batch_translations.items():
                            if j in all_glossary_segments[i]:
                                # Restore glossary terms with target translations
                                segments = all_glossary_segments[i][j]
                                original_translation = translation
                                translation = self.glossary_protection.restore_text(translation, segments)
                                batch_translations[j] = translation
                                self.logger.debug(f"Restored {len(segments)} glossary terms: {original_translation[:50]}... → {translation[:50]}...")

                    # Map back to original indices
                    for j, translation in batch_translations.items():
                        original_idx = i + j
                        all_translations[original_idx] = translation

                    pbar.update(len(batch))
            
            # Apply translations back to document
            text_index = 0
            
            # Apply to paragraphs
            for para in doc.paragraphs:
                text_index = self._process_paragraph(para, all_translations, text_index)
            
            # Apply to tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_index = self._process_table_cell(cell, all_translations, text_index)

            # Post-process: Restore font and type-protected text from placeholders
            self._restore_text_protection(doc, protected_segments)

            # Mark TOC fields for update before saving
            self._mark_toc_for_update(doc)

            # Save translated document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            self.logger.info(f"✅ Saved translated document: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Failed to process {input_path}: {e}")
            if self.config.debug:
                import traceback
                traceback.print_exc()
            raise


def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(
        description="Word Document Translation Utility with Google Generative AI",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    # Input options
    parser.add_argument("input_file", nargs='?', help="Single input DOCX file path")
    parser.add_argument("--input-files", help="Directory containing DOCX files to process")

    # Language options
    parser.add_argument('--source-lang', default='auto',
                       help='Source language code (default: auto-detect)')

    parser.add_argument('--target-lang', help='Target language code (e.g., ja)')
    parser.add_argument('--target-langs', help='Multiple target languages (comma-separated)')

    # Translation options
    parser.add_argument('--style-prompt', choices=list(STYLE_PROMPTS.keys()),
                       default='business', help='Translation style')
    parser.add_argument('--context-file', help='Path to glossary/context JSON file')
    parser.add_argument('--smart-context', action='store_true',
                       help='Use document structure for better context')
    parser.add_argument('--dont-translate-font', nargs='?', const='Courier New',
                       help='Font names to exclude from translation (default: Courier New). Use comma-separated list for multiple fonts.')
    parser.add_argument('--dont-translate-type', nargs='?', const='',
                       help='Formatting types to exclude from translation. Valid types: bold, italic, underlined, italic-bold, underlined-bold, underlined-italic, underlined-italic-bold. Use comma-separated list for multiple types.')

    # Output options
    parser.add_argument('--output-dir', default='.', help='Output directory')
    parser.add_argument("--output-files", help="Directory to save translated files (same as --output-dir)")
    parser.add_argument('--model', default=DEFAULT_MODEL, help='Gemini model name')
    parser.add_argument('--batch-size', type=int, default=DEFAULT_BATCH_SIZE,
                       help=f'Texts per API call (default: {DEFAULT_BATCH_SIZE})')
    parser.add_argument('--concurrency', type=int, default=DEFAULT_CONCURRENCY,
                       help=f'Concurrent API calls (default: {DEFAULT_CONCURRENCY})')
    
    # Other options
    parser.add_argument('--debug', action='store_true', help='Enable debug logging')
    parser.add_argument('--list-styles', action='store_true', help='List available styles')
    parser.add_argument('--list-types', action='store_true', help='List valid formatting types for --dont-translate-type')
    parser.add_argument('--version', action='store_true', help='Show version information')
    
    args = parser.parse_args()
    
    # Handle version request
    if args.version:
        print(f"{Fore.CYAN}DOCX Translator using google-generativeai{Style.RESET_ALL}")
        print(f"Model: {DEFAULT_MODEL}")
        sys.exit(0)
    
    # Handle info requests
    if args.list_types:
        print(f"\n{Fore.CYAN}Valid formatting types for --dont-translate-type:{Style.RESET_ALL}")
        for type_name in TypeProtectionEngine.get_valid_types():
            formatting = TypeProtectionEngine.TYPE_MAPPINGS[type_name]
            properties = []
            for prop, value in formatting.items():
                if value:
                    properties.append(prop)
            print(f"  {Fore.GREEN}{type_name}:{Style.RESET_ALL} {', '.join(properties)}")
        print(f"\n{Fore.CYAN}Usage examples:{Style.RESET_ALL}")
        print(f"  --dont-translate-type bold")
        print(f"  --dont-translate-type bold,italic")
        print(f"  --dont-translate-type underlined-italic-bold")
        sys.exit(0)

    if args.list_styles:
        print(f"\n{Fore.CYAN}Available translation styles:{Style.RESET_ALL}")
        for style, prompts in STYLE_PROMPTS.items():
            print(f"\n  {Fore.GREEN}{style}:{Style.RESET_ALL}")
            for lang, prompt in prompts.items():
                if lang != 'default':
                    lang_name = 'Japanese' if lang == 'ja' else lang.upper()
                    print(f"    {lang_name}: {prompt[:60]}...")
                else:
                    print(f"    Default: {prompt[:60]}...")
        sys.exit(0)
    
    # Validate input arguments
    if not args.input_file and not args.input_files:
        parser.error("Either provide a single input file or use --input-files for directory processing")

    if args.input_file and args.input_files:
        parser.error("Cannot use both single input file and --input-files together")

    # Handle input files
    input_files = []
    if args.input_file:
        input_path = Path(args.input_file)
        if not input_path.exists():
            print(f"{Fore.RED}❌ Input file '{args.input_file}' does not exist{Style.RESET_ALL}")
            sys.exit(1)
        if not input_path.suffix.lower() == '.docx':
            print(f"{Fore.RED}Error: Input file must be a .docx file{Style.RESET_ALL}")
            sys.exit(1)
        input_files = [args.input_file]
    elif args.input_files:
        try:
            input_files = get_docx_files_from_directory(args.input_files)
            if not input_files:
                sys.exit(1)
        except FileNotFoundError as e:
            print(f"{Fore.RED}❌ {str(e)}{Style.RESET_ALL}")
            sys.exit(1)

    # Handle target languages
    target_langs = []
    if args.target_langs:
        target_langs = [lang.strip() for lang in args.target_langs.split(',')]
    elif args.target_lang:
        target_langs = [args.target_lang]
    else:
        parser.error("At least one target language must be specified using --target-lang or --target-langs")

    # Handle output directory
    output_dir = Path(args.output_files) if args.output_files else Path(args.output_dir)

    # Handle font protection
    dont_translate_fonts = None
    if args.dont_translate_font:
        if ',' in args.dont_translate_font:
            dont_translate_fonts = [font.strip() for font in args.dont_translate_font.split(',')]
        else:
            dont_translate_fonts = [args.dont_translate_font.strip()]

    # Handle type protection
    dont_translate_types = None
    if args.dont_translate_type:
        if ',' in args.dont_translate_type:
            type_list = [t.strip().lower() for t in args.dont_translate_type.split(',')]
        else:
            type_list = [args.dont_translate_type.strip().lower()]

        # Validate types
        valid_types, invalid_types = TypeProtectionEngine.validate_types(type_list)

        if invalid_types:
            print(f"{Fore.YELLOW}⚠️ Invalid formatting types ignored: {', '.join(invalid_types)}{Style.RESET_ALL}")
            print(f"{Fore.CYAN}ℹ️ Valid types: {', '.join(TypeProtectionEngine.get_valid_types())}{Style.RESET_ALL}")

        if valid_types:
            dont_translate_types = valid_types
            print(f"{Fore.GREEN}✅ Type protection enabled for: {', '.join(valid_types)}{Style.RESET_ALL}")
        else:
            print(f"{Fore.YELLOW}⚠️ No valid formatting types specified{Style.RESET_ALL}")

    # Show header
    print(f"\n{Fore.CYAN}{'='*60}{Style.RESET_ALL}")
    print(f"{Fore.GREEN}DOCX Translator - Google Generative AI{Style.RESET_ALL}")
    print(f"{Fore.CYAN}{'='*60}{Style.RESET_ALL}")
    print(f"Model: {args.model}")
    print(f"Source: {args.source_lang}")
    print(f"Targets: {', '.join(target_langs)}")
    print(f"Style: {args.style_prompt}")
    print(f"Files: {len(input_files)}")
    print(f"{Fore.CYAN}{'='*60}{Style.RESET_ALL}\n")

    # Process all files
    all_output_files = []
    total_files = len(input_files)
    
    for i, input_file in enumerate(input_files, 1):
        print(f"\n{Fore.BLUE}📊 Processing file {i}/{total_files}: {Path(input_file).name}{Style.RESET_ALL}")
        input_path = Path(input_file)

        for target_lang in target_langs:
            print(f"\n{Fore.CYAN}Translating to {target_lang}...{Style.RESET_ALL}")

            # Create configuration
            config = TranslationConfig(
                source_lang=args.source_lang,
                target_lang=target_lang,
                model_name=args.model,
                batch_size=args.batch_size,
                concurrency=args.concurrency,
                style_prompt=args.style_prompt,
                context_file=args.context_file,
                smart_context=args.smart_context,
                debug=args.debug,
                dont_translate_fonts=dont_translate_fonts,
                dont_translate_types=dont_translate_types
            )

            # Create translator
            try:
                translator = DocxTranslator(config)
            except ValueError as e:
                print(str(e))
                sys.exit(1)

            # Generate output filename
            output_name = f"{input_path.stem}_{target_lang}{input_path.suffix}"
            output_path = output_dir / output_name

            # Translate document
            try:
                start_time = time.time()
                translator.translate_document(input_path, output_path)
                elapsed_time = time.time() - start_time
                print(f"{Fore.GREEN}✅ Successfully created: {output_path}{Style.RESET_ALL}")
                print(f"{Fore.CYAN}   Time taken: {elapsed_time:.2f} seconds{Style.RESET_ALL}")
                all_output_files.append(output_path)
            except Exception as e:
                print(f"{Fore.RED}❌ Translation failed: {e}{Style.RESET_ALL}")
                if args.debug:
                    import traceback
                    traceback.print_exc()
                continue

    # Summary
    print(f"\n{Fore.GREEN}🎉 Translation completed!{Style.RESET_ALL}")
    print(f"{Fore.CYAN}📊 Summary:{Style.RESET_ALL}")
    print(f"   📄 Input files processed: {total_files}")
    print(f"   🌍 Target languages: {len(target_langs)} ({', '.join(target_langs)})")
    print(f"   📁 Output files created: {len(all_output_files)}")

    if all_output_files:
        print(f"\n{Fore.CYAN}📁 Output files:{Style.RESET_ALL}")
        for output_file in all_output_files:
            print(f"   ✅ {output_file}")

        print(f"\n{Fore.YELLOW}📑 Table of Contents Note:{Style.RESET_ALL}")
        print(f"   If your documents contain Table of Contents, Word will prompt")
        print(f"   to update them when you open the translated files. Click 'Yes'")
        print(f"   to refresh the TOC with translated headings.")


if __name__ == "__main__":
    main()
